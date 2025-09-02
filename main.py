
import os
import asyncio
from typing import Dict, List, Any, Annotated, TypedDict
from datetime import datetime

# Environment and configuration
from dotenv import load_dotenv

# LangChain core components
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.tools import tool
from langchain_google_genai import ChatGoogleGenerativeAI

# LangGraph components
from langgraph.graph import StateGraph, END, START
from langgraph.graph.message import add_messages

# Tavily for research
from tavily import TavilyClient

# Document creation
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Load environment variables
load_dotenv()

# Define the state
class ResearchState(TypedDict):
    messages: Annotated[list, add_messages]
    topic: str
    research_data: str
    paper_content: str
    document_path: str
    pdf_path: str
    step: str

# Standalone tool functions
def research_with_tavily(topic: str) -> str:
    """Research using Tavily API"""
    try:
        tavily_client = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))
        
        search_results = tavily_client.search(
            query=f"academic research {topic} recent studies findings",
            search_depth="advanced",
            max_results=6
        )
        
        formatted_results = f"Research Results for '{topic}':\n\n"
        for i, result in enumerate(search_results.get('results', []), 1):
            formatted_results += f"{i}. **{result.get('title', '')}**\n"
            formatted_results += f"   Content: {result.get('content', '')[:600]}...\n"
            formatted_results += f"   Source: {result.get('url', '')}\n\n"
        
        return formatted_results
        
    except Exception as e:
        return f"Research failed: {str(e)}. Using AI knowledge instead."

def create_documents(content: str, topic: str) -> Dict[str, str]:
    """Create Word and PDF documents"""
    try:
        # Ensure doc directory exists
        doc_dir = os.path.join(os.getcwd(), "doc")
        os.makedirs(doc_dir, exist_ok=True)
        
        # Create filename
        safe_topic = "".join(c if c.isalnum() or c in (' ', '-', '_') else '' for c in topic)
        safe_topic = safe_topic.replace(' ', '_')[:30]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create Word document
        filename = f"research_paper_{safe_topic}_{timestamp}.docx"
        document_path = os.path.join(doc_dir, filename)
        
        doc = DocxDocument()
        
        # Parse content and add to document
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            else:
                if line and not line.startswith('#'):
                    doc.add_paragraph(line)
        
        doc.save(document_path)
        
        # Create PDF
        pdf_path = document_path.replace('.docx', '.pdf')
        try:
            doc_pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1
            )
            
            # Add content to PDF
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 12))
                    continue
                
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], title_style))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                else:
                    if line and not line.startswith('#'):
                        story.append(Paragraph(line, styles['Normal']))
                        story.append(Spacer(1, 6))
            
            doc_pdf.build(story)
            
        except Exception as e:
            pdf_path = f"PDF creation failed: {e}"
        
        return {
            'document_path': document_path,
            'pdf_path': pdf_path
        }
        
    except Exception as e:
        return {
            'document_path': f"Document creation failed: {e}",
            'pdf_path': f"PDF creation failed: {e}"
        }

class SimpleResearchAgent:
    def __init__(self):
        """Initialize the research agent"""
        
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            google_api_key=os.getenv("GOOGLE_API_KEY"),
            temperature=0.7
        )
        
        self.workflow = self._create_workflow()
        self.app = self.workflow.compile()
    
    def _create_workflow(self) -> StateGraph:
        """Create a simple linear workflow"""
        
        workflow = StateGraph(ResearchState)
        
        # Add nodes
        workflow.add_node("research", self._research_node)
        workflow.add_node("write", self._write_node)
        workflow.add_node("create_docs", self._create_docs_node)
        
        # Simple linear flow
        workflow.add_edge(START, "research")
        workflow.add_edge("research", "write")
        workflow.add_edge("write", "create_docs")
        workflow.add_edge("create_docs", END)
        
        return workflow
    
    def _research_node(self, state: ResearchState) -> ResearchState:
        """Research node"""
        
        topic = state.get("topic", "")
        print(f"🔍 Researching: {topic}")
        
        # Do research
        research_data = research_with_tavily(topic)
        
        return {
            **state,
            "research_data": research_data,
            "step": "research_complete"
        }
    
    def _write_node(self, state: ResearchState) -> ResearchState:
        """Writing node"""
        
        topic = state.get("topic", "")
        research_data = state.get("research_data", "")
        
        print(f"✍️ Writing paper for: {topic}")
        
        # Create writing prompt
        writing_prompt = f"""Based on the research data below, write a comprehensive academic research paper about "{topic}".

Research Data:
{research_data}

Structure the paper with:
1. # Title
2. ## Abstract (150-200 words)
3. ## Introduction
4. ## Literature Review/Background
5. ## Main Analysis (2-3 sections with ### subheadings)
6. ## Conclusion
7. ## References

Use formal academic language and ensure the paper is well-researched and comprehensive.
Format using markdown headers (# ## ###).
Include proper citations and references.
Aim for approximately 2000-2500 words.

Write the complete paper now:"""
        
        # Generate the paper
        response = self.llm.invoke([HumanMessage(content=writing_prompt)])
        paper_content = response.content
        
        return {
            **state,
            "paper_content": paper_content,
            "step": "writing_complete"
        }
    
    def _create_docs_node(self, state: ResearchState) -> ResearchState:
        """Document creation node"""
        
        topic = state.get("topic", "")
        paper_content = state.get("paper_content", "")
        
        print(f"📄 Creating documents for: {topic}")
        
        # Create documents
        doc_results = create_documents(paper_content, topic)
        
        return {
            **state,
            "document_path": doc_results['document_path'],
            "pdf_path": doc_results['pdf_path'],
            "step": "documents_complete"
        }
    
    async def create_research_paper(self, topic: str) -> Dict[str, Any]:
        """Create a research paper using the simplified workflow"""
        
        print(f"🚀 Starting research paper creation for: {topic}")
        
        initial_state = {
            "messages": [],
            "topic": topic,
            "research_data": "",
            "paper_content": "",
            "document_path": "",
            "pdf_path": "",
            "step": "start"
        }
        
        try:
            # Run the workflow
            final_state = await asyncio.to_thread(self.app.invoke, initial_state)
            
            return {
                "topic": topic,
                "document_path": final_state.get("document_path", ""),
                "pdf_path": final_state.get("pdf_path", ""),
                "paper_content": final_state.get("paper_content", ""),
                "status": "completed" if final_state.get("step") == "documents_complete" else "partial"
            }
            
        except Exception as e:
            print(f"❌ Error in workflow: {e}")
            return {
                "topic": topic,
                "error": str(e),
                "status": "failed"
            }
    
    async def edit_paper(self, current_content: str, change_request: str, topic: str) -> Dict[str, Any]:
        """Edit the research paper based on user request"""
        
        editing_prompt = f"""You are editing a research paper. The user wants to make the following changes:

USER REQUEST: {change_request}

CURRENT PAPER CONTENT:
{current_content}

Please apply the requested changes to the paper content. Maintain the same academic structure and formatting, but incorporate the user's specific requests. Return the complete modified paper content with proper markdown formatting.

Modified paper:"""
        
        
            # Generate the edited paper
        response = self.llm.invoke([HumanMessage(content=editing_prompt)])
        edited_content = response.content
            
            # Create new documents with the edited content
        doc_results = create_documents(edited_content, topic)
            
        return{
                "topic": topic,
                "document_path": doc_results['document_path'],
                "pdf_path": doc_results['pdf_path'],
                "paper_content": edited_content,
                "status": "completed"
            }
            
        


def main():
    """Main function"""
    
    print("🎓 AI Research Paper Creator")
    print("=" * 40)
    
  
    # Get user input
    topic = input("\nEnter your research topic: ").strip()
   
    
    # Create the agent and start processing immediately
    agent = SimpleResearchAgent()
    
    # Create the research paper
    try:
        print(f"\n🔄 Creating research paper on: {topic}")
        print("🔍 Starting research and writing process...")
        
        result = asyncio.run(agent.create_research_paper(topic))
        
        print("\n" + "=" * 50)
        print("📊 RESEARCH PAPER COMPLETED")
        print("=" * 50)
        
        if result.get("status") == "completed":
            print(f"✅ Topic: {result['topic']}")
            print(f"📄 Word Document: {result['document_path']}")
            print(f"📑 PDF Document: {result['pdf_path']}")
            
            print(f"\n📁 Documents saved in: doc/ folder")
            
            # Store the current result for potential editing
            current_result = result
            
            # Ask for changes
            while True:
                print("\n" + "=" * 50)
                make_changes = input("\nWould you like to make any changes to the paper? (y/n): ").lower().strip()
                
                if make_changes == 'y':
                    change_request = input("\nDescribe the changes you want to make: ").strip()
                    if change_request:
                        print(f"\n🔄 Applying changes: {change_request}")
                        
                        # Apply the changes
                        try:
                            edit_result = asyncio.run(agent.edit_paper(
                                current_result.get("paper_content", ""), 
                                change_request, 
                                current_result.get("topic", "")
                            ))
                            
                            if edit_result.get("status") == "completed":
                                print("✅ Changes applied successfully!")
                                print(f"📄 Updated Word Document: {edit_result['document_path']}")
                                print(f"📑 Updated PDF Document: {edit_result['pdf_path']}")
                                
                                # Update current result for further editing
                                current_result = edit_result
                            else:
                                print(f"❌ Edit failed: {edit_result.get('error', 'Unknown error')}")
                                
                        except Exception as e:
                            print(f"❌ Edit error: {e}")
                        
                    else:
                        print("❌ Please provide a valid change request.")
                elif make_changes == 'n':
                    print("\n✅ Research paper creation completed successfully!")
                    break
                else:
                    print("❌ Please enter 'y' for yes or 'n' for no.")
            
        else:
            print(f"❌ Status: {result.get('status', 'failed').upper()}")
            print(f"Error: {result.get('error', 'Unknown error')}")
            
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    main()
